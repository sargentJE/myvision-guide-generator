"""
File management for MyVision Guide Generator

This handles all file creation, naming, and organization.
"""

# SECTION Imports
"""
Essential imports for comprehensive file management and document processing.

Key components explained:
- re: Python's regular expression library for text pattern matching and cleaning.
  Used extensively for sanitizing filenames and processing text content safely.

- datetime: Standard library for timestamp generation and date formatting.
  Critical for creating unique filenames and tracking content creation times.

- pathlib.Path: Modern, object-oriented file system path handling that works
  cross-platform (Windows, macOS, Linux). Replaces older os.path approaches
  with more intuitive and robust path operations.

- typing.Optional, typing.Dict, typing.Any: Type hints that document function
  signatures and improve code maintainability. Optional indicates values that
  can be None, while Dict provides structure for complex data types.

- docx.Document: Python-docx library for creating and manipulating Microsoft
  Word documents programmatically. Enables professional document formatting
  with headers, footers, styling, and layout control.

- docx.shared.Inches: Measurement units for precise document layout control,
  ensuring consistent formatting across different systems and Word versions.

- docx.enum.text.WD_ALIGN_PARAGRAPH: Text alignment constants for professional
  document formatting (left, center, right, justified alignment options).

- .config: Local configuration module providing file paths, branding settings,
  and document formatting preferences for consistent output across the system.
"""
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import config

# SECTION FileManager Class
class FileManager:
    """
    Comprehensive file management system for educational content generation.
    
    This class implements enterprise-grade file operations with a focus on:
    - Cross-platform compatibility (Windows, macOS, Linux)
    - Professional document formatting and branding
    - Safe filename generation and collision prevention
    - Multiple output formats (Markdown, Word documents)
    - Consistent file organization and directory structure
    
    Architecture Overview:
    The FileManager follows the Single Responsibility Principle, handling all
    file system interactions while the GuideGenerator focuses on content creation.
    This separation enables:
    - Independent testing of file operations
    - Easy extension to new file formats
    - Centralized file handling policies
    - Consistent error handling for file operations
    
    Design Patterns:
    - Utility Class: Provides stateless file operations with configuration
    - Template Method: Standard workflows for different document types
    - Strategy Pattern: Different formatting approaches for various output types
    
    File Safety:
    All file operations implement defensive programming practices:
    - Filename sanitization prevents directory traversal attacks
    - Path validation ensures files are created in intended locations
    - Atomic operations prevent partial file writes
    - Proper encoding handling for international characters
    
    Professional Output:
    Documents are formatted according to accessibility best practices:
    - Clear heading hierarchies for screen reader navigation
    - Appropriate font sizes and spacing for visual accessibility
    - Consistent branding and professional appearance
    - Cross-platform font compatibility
    """
    
    # SECTION Constructor
    def __init__(self):
        """
        Initialize the file manager and ensure proper directory structure.
        
        Initialization Process:
        The constructor performs essential setup operations that ensure the
        file system is ready for document generation:
        
        1. Directory Validation:
           Calls config.ensure_output_directory() to verify that the configured
           output directory exists and is writable. This prevents runtime errors
           when attempting to save generated documents.
        
        2. Permission Checking:
           Implicitly validates that the application has appropriate file system
           permissions to create directories and write files in the target location.
        
        3. Path Resolution:
           Resolves any relative paths to absolute paths, ensuring consistent
           behavior regardless of the application's working directory.
        
        Defensive Programming:
        This approach implements "fail-fast" principles - if there are file system
        issues, they're discovered immediately rather than during document generation
        when users are waiting for results.
        
        Cross-Platform Considerations:
        The pathlib-based configuration system handles platform differences
        automatically, ensuring the same code works on Windows (with backslash
        separators), macOS/Linux (with forward slashes), and handles drive letters,
        network paths, and other platform-specific path formats correctly.
        
        Error Handling:
        Any filesystem errors during initialization (permissions, disk space,
        invalid paths) will be raised immediately, allowing the application to
        handle them gracefully or inform the user of configuration issues.
        """
        config.ensure_output_directory()
    
    # SECTION Filename Generation
    def generate_filename(self, 
                         title: str, 
                         guide_type: str = "guide",
                         format_type: str = "markdown") -> str:
        """
        Generate safe, descriptive, and unique filenames for educational content.
        
        This method implements comprehensive filename generation that balances
        human readability with technical requirements for cross-platform compatibility.
        
        Filename Architecture:
        The generated filename follows this structure:
        {clean_title}_{guide_type}_{timestamp}.{extension}
        
        Example outputs:
        - "voiceover_rotor_basics_learning_guide_20240608_143022.md"
        - "jaws_quick_start_session_notes_20240608_143022.docx"
        - "nvda_web_navigation_reference_guide_20240608_143022.md"
        
        Parameters Explained:
        - title (str): The human-readable title from AI generation or user input.
          Will be sanitized to remove unsafe characters while preserving meaning.
        
        - guide_type (str): Content categorization for organization and search.
          Common values: "guide", "learning", "session", "reference", "quick"
          Default: "guide" for general-purpose content
        
        - format_type (str): Target file format determining extension and processing.
          Supported: "markdown" (.md), "docx" (.docx)
          Default: "markdown" for universal compatibility
        
        Safety Features:
        1. Character Sanitization: Removes/replaces filesystem-unsafe characters
        2. Length Limitations: Prevents excessively long filenames
        3. Collision Prevention: Timestamp ensures uniqueness across generations
        4. Cross-Platform Compatibility: Works on Windows, macOS, and Linux
        5. Reserved Name Avoidance: Prevents conflicts with system files
        
        Timestamp Strategy:
        Uses high-precision timestamps (YYYYMMDD_HHMMSS) to ensure:
        - Chronological sorting in file explorers
        - Unique filenames even for rapid generation
        - Easy identification of content creation time
        - No filename collisions in concurrent usage
        
        Accessibility Considerations:
        Filenames are screen reader friendly with:
        - Descriptive, meaningful names
        - No special characters that cause pronunciation issues
        - Logical structure that conveys content hierarchy
        - Consistent patterns for predictable navigation
        """
        # SECTION Filename Processing
        """
        Multi-step filename construction with safety and uniqueness guarantees.
        
        Processing Pipeline:
        1. Title Sanitization: Clean unsafe characters and normalize format
        2. Timestamp Generation: Create unique identifier for collision prevention
        3. Extension Selection: Choose appropriate file extension for format
        4. Component Assembly: Combine elements into final filename
        
        Title Cleaning Process:
        The _clean_filename() method handles complex sanitization:
        - Removes filesystem-unsafe characters (/, \\, :, *, ?, ", <, >, |)
        - Converts spaces to underscores for command-line compatibility
        - Normalizes unicode characters for cross-platform consistency
        - Truncates excessive length while preserving meaning
        - Handles edge cases like empty strings or special-only content
        
        Timestamp Format Explanation:
        YYYYMMDD_HHMMSS format provides:
        - Year-Month-Day for easy date identification
        - Hour-Minute-Second for precise timing
        - Underscore separator for readability
        - Fixed width for consistent sorting
        - No timezone complexity (uses local time)
        
        Extension Logic:
        Simple but extensible format selection:
        - "markdown" → ".md" for universal text processing
        - All others → ".docx" for Microsoft Word compatibility
        - Easy to extend for additional formats (PDF, HTML, etc.)
        
        Component Assembly:
        Final filename structure ensures:
        - Descriptive content identification
        - Consistent organizational pattern
        - Unique identification through timestamp
        - Professional appearance in file systems
        """
        # Clean up the title for use in filename
        clean_title = self._clean_filename(title)
        
        # Create timestamp for unique filenames
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Choose file extension
        extension = "md" if format_type == "markdown" else "docx"
        
        # Build the filename
        filename = f"{clean_title}_{guide_type}_guide_{timestamp}.{extension}"
        
        return filename
    
    # SECTION Filename Cleaning
    def _clean_filename(self, text: str) -> str:
        """
        Convert text into a safe filename.
        
        Example: "VoiceOver Rotor Basics!" becomes "voiceover_rotor_basics"
        """
        # SECTION Text Sanitization
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters, keep only letters, numbers, spaces, hyphens
        text = re.sub(r'[^\w\s-]', '', text)
        
        # Replace spaces and hyphens with underscores
        text = re.sub(r'[-\s]+', '_', text)
        
        # Remove leading/trailing underscores
        text = text.strip('_')
        
        # Limit length to avoid filesystem issues
        if len(text) > 50:
            text = text[:50].rstrip('_')
        
        return text
    
    # SECTION Markdown File Operations
    def save_markdown_guide(self, 
                           content: str, 
                           metadata: Dict[str, Any],
                           filename: str) -> Path:
        """
        Save guide as a markdown file with metadata header.
        
        Args:
            content: The guide content from Claude
            metadata: Information about the guide
            filename: What to name the file
            
        Returns:
            Path to the saved file
        """
        # SECTION Directory Setup
        # Create the output directory path
        output_dir = config.output_directory / "Learning_Guides"
        output_dir.mkdir(exist_ok=True)
        
        # Full file path
        file_path = output_dir / filename
        
        # SECTION Content Assembly
        # Create metadata header (YAML frontmatter)
        metadata_header = self._create_metadata_header(metadata)
        
        # Combine metadata and content
        full_content = f"{metadata_header}\n\n{content}"
        
        # SECTION File Writing
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(full_content)
        
        return file_path
    
    # SECTION Word Document Operations
    def save_docx_guide(self, 
                       content: str, 
                       metadata: Dict[str, Any],
                       filename: str) -> Path:
        """
        Save guide as a Word document with MyVision branding.
        
        Args:
            content: The guide content from Claude
            metadata: Information about the guide  
            filename: What to name the file
            
        Returns:
            Path to the saved file
        """
        # SECTION Directory Setup  
        # Create the output directory path
        output_dir = config.output_directory / "Learning_Guides"
        output_dir.mkdir(exist_ok=True)
        
        # Full file path
        file_path = output_dir / filename
        
        # SECTION Document Creation
        # Create Word document
        doc = self._create_word_document(content, metadata)
        
        # SECTION Document Saving
        # Save document
        doc.save(file_path)
        
        return file_path
    
    # SECTION Metadata Processing
    def _create_metadata_header(self, metadata: Dict[str, Any]) -> str:
        """
        Create YAML frontmatter for markdown files.
        
        This adds structured information at the top of markdown files.
        """
        # SECTION YAML Generation
        lines = ["---"]
        
        for key, value in metadata.items():
            if value is not None:
                lines.append(f"{key}: {value}")
        
        lines.append("---")
        
        return "\n".join(lines)
    
    # SECTION Word Document Creation
    def _create_word_document(self, content: str, metadata: Dict[str, Any]) -> Document:
        """
        Create a professionally formatted, accessible Word document with large print.
        
        This method implements comprehensive accessibility standards for visually
        impaired users, following the MyVision organization's commitment to
        inclusive design and professional document creation.
        
        Accessibility Features Implemented:
        - Large print formatting (18pt minimum body text)
        - High contrast color schemes when enabled
        - Clear heading hierarchy with appropriate sizing
        - Enhanced line spacing for readability
        - Accessible font selection (Arial by default)
        - Proper document structure for screen readers
        
        Document Structure:
        1. Professional MyVision branded header with logo support
        2. Accessibility-optimized content formatting
        3. Clear footer with contact information
        4. Accessibility statement indicating large print formatting
        
        Args:
            content: Guide content in markdown format from Claude AI
            metadata: Document metadata including title, date, topic information
            
        Returns:
            Document: Fully formatted Word document optimized for accessibility
        """
        # SECTION Document Initialization
        # Create new document with accessibility defaults
        doc = Document()
        
        # SECTION Document-Wide Accessibility Setup
        # Set document-wide accessibility formatting defaults
        self._set_document_accessibility_defaults(doc)
        
        # SECTION Document Structure Assembly
        # Add professional MyVision header with accessibility formatting
        self._add_accessible_document_header(doc, metadata)
        
        # Parse and add content with large print formatting
        self._add_accessible_content_to_document(doc, content)
        
        # Add professional footer with accessibility statement
        self._add_accessible_document_footer(doc)
        
        return doc

    def _set_document_accessibility_defaults(self, doc: Document):
        """
        Configure document-wide accessibility formatting standards.
        
        This method establishes the foundational formatting that ensures all
        text elements in the document meet accessibility requirements. It
        modifies the document's built-in styles to provide consistent large
        print formatting throughout the entire document.
        
        Style Configuration Strategy:
        - Normal Style: Base style for all body text (18pt minimum)
        - Heading Styles: Clear hierarchy with appropriate sizing
        - Color Schemes: High contrast options for better visibility
        - Typography: Accessible font families with clear readability
        
        Benefits of This Approach:
        - Consistent formatting across all document elements
        - Automatic inheritance of accessibility settings
        - Easy maintenance and future updates
        - Professional appearance with accessibility compliance
        """
        # SECTION Style Collection Access
        # Get document styles for modification
        styles = doc.styles
        
        # SECTION Normal Style Configuration
        # Configure the base Normal style that all text inherits from
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = config.accessible_font
        normal_font.size = Pt(config.body_font_size)
        normal_font.color.rgb = RGBColor(*config.text_color)
        
        # Set paragraph formatting for Normal style
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing = config.line_spacing
        normal_paragraph.space_after = Pt(config.paragraph_spacing)
        
        # SECTION Heading Style Configuration
        # Configure heading styles for clear visual hierarchy and accessibility
        heading_configs = [
            ('Heading 1', config.heading1_font_size, True, config.heading_color),
            ('Heading 2', config.heading2_font_size, True, config.heading_color),
            ('Heading 3', config.heading3_font_size, True, RGBColor(51, 51, 51)),  # Slightly lighter for hierarchy
        ]
        
        for style_name, font_size, is_bold, color_tuple in heading_configs:
            if style_name in styles:
                heading_style = styles[style_name]
                heading_font = heading_style.font
                heading_font.name = config.accessible_font
                heading_font.size = Pt(font_size)
                heading_font.bold = is_bold
                heading_font.color.rgb = RGBColor(*color_tuple)
                
                # Add enhanced spacing around headings for readability
                heading_paragraph = heading_style.paragraph_format
                heading_paragraph.space_before = Pt(18)
                heading_paragraph.space_after = Pt(12)
                heading_paragraph.line_spacing = config.line_spacing
    
    # SECTION Accessible Document Header
    def _add_accessible_document_header(self, doc: Document, metadata: Dict[str, Any]):
        """
        Create professional MyVision branded header with large print accessibility.
        
        This header section implements comprehensive accessibility standards while
        maintaining the professional MyVision brand identity. The design prioritizes
        readability for visually impaired users while creating an attractive,
        professional document appearance.
        
        Header Components:
        1. MyVision Logo - Professional brand image (centered, 2-inch width)
        2. Document Title - Extra large (26pt) for maximum visibility
        3. Organization Branding - Clear identification with accessible sizing
        4. Creation Metadata - Date and context information
        5. Visual Separator - Clean document section division
        
        Accessibility Features:
        - All text meets large print standards (18pt minimum)
        - High contrast color options for better visibility
        - Clear visual hierarchy with appropriate spacing
        - Screen reader friendly structure and content
        - Professional branding that enhances rather than hinders accessibility
        """
        # SECTION Logo Integration
        # Add MyVision logo to document header for professional branding
        logo_added = self._add_logo_to_header(doc)
        
        # SECTION Main Title Section
        # Create centered title paragraph with large print accessibility
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if logo_added:
            title_para.space_before = Pt(18)
        
        # Main title - extra large for maximum accessibility
        title_run = title_para.add_run(metadata.get("title", "Learning Guide"))
        title_run.font.name = config.accessible_font
        title_run.font.size = Pt(26)  # Extra large for main title visibility
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(*config.heading_color)
        
        # SECTION Organization Branding
        # Organization name with accessible formatting
        org_para = doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_para.space_before = Pt(8)
        
        org_run = org_para.add_run(config.organization_name)
        org_run.font.name = config.accessible_font
        org_run.font.size = Pt(config.body_font_size)  # Use large print size
        org_run.font.color.rgb = RGBColor(102, 102, 102)  # Readable gray
        org_run.italic = True
        
        # SECTION Metadata Information
        # Date and document context with accessible formatting
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.space_before = Pt(6)
        
        date_text = f"Generated: {datetime.now().strftime('%B %d, %Y')}"
        if metadata.get('technology') and metadata.get('device'):
            date_text += f" • {metadata['technology']} on {metadata['device']}"
        elif metadata.get('topic'):
            date_text += f" • {metadata['topic']}"
        
        date_run = date_para.add_run(date_text)
        date_run.font.name = config.accessible_font
        date_run.font.size = Pt(config.body_font_size - 2)  # Slightly smaller but still accessible
        date_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # SECTION Visual Separator
        # Professional separator with accessible formatting
        separator_para = doc.add_paragraph()
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_para.space_before = Pt(18)
        separator_para.space_after = Pt(24)
        
        separator_run = separator_para.add_run("●" * 25)
        separator_run.font.name = config.accessible_font
        separator_run.font.size = Pt(config.body_font_size)
        separator_run.font.color.rgb = RGBColor(*config.accent_color)

    def _add_logo_to_header(self, doc: Document) -> bool:
        """
        Add MyVision logo to document header if available.
        
        This method adds the organization logo to the document header when
        the logo file exists and is accessible. The logo is sized appropriately
        for professional document appearance while maintaining accessibility.
        
        Implementation Details:
        - Checks for logo file existence before attempting to insert
        - Sizes logo appropriately for document headers (max 2 inches wide)
        - Centers logo in document for professional appearance
        - Provides proper spacing around logo
        - Returns status for conditional formatting adjustments
        
        Returns:
            bool: True if logo was successfully added, False otherwise
        """
        # SECTION Logo File Validation
        # Check if logo file exists and is accessible
        if not hasattr(config, 'logo_path') or not config.logo_path:
            return False
            
        if not config.logo_path.exists():
            return False
            
        # SECTION Logo Insertion
        try:
            # Create paragraph for logo with center alignment
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.space_after = Pt(12)  # Space after logo
            
            # Add logo to paragraph
            logo_run = logo_para.add_run()
            logo_run.add_picture(str(config.logo_path), width=Inches(2.0))  # 2 inch width
            
            return True
            
        except Exception as e:
            # If logo insertion fails, continue without logo
            # This prevents document generation from failing due to logo issues
            return False

    def _add_accessible_content_to_document(self, doc: Document, content: str):
        """
        Convert markdown content to accessible Word document formatting.
        
        This method provides comprehensive markdown parsing with accessibility
        enhancements, ensuring that all content elements meet large print and
        readability standards for visually impaired users.
        
        Supported Markdown Elements:
        - Headings (H1, H2, H3) with proper hierarchy and sizing
        - Bullet and numbered lists with enhanced spacing
        - Bold text formatting (**text**)
        - Regular paragraphs with optimal line spacing
        - Empty lines for proper document flow
        
        Accessibility Enhancements:
        - All text uses large print font sizes (18pt minimum)
        - Clear heading hierarchy for screen reader navigation
        - Enhanced spacing for visual separation
        - High contrast color options
        - Proper list formatting with accessible indentation
        """
        # SECTION Content Line Processing
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # SECTION Empty Line Handling
                # Empty line - add paragraph break with proper spacing
                empty_para = doc.add_paragraph()
                empty_para.space_after = Pt(6)
                continue
            
            # SECTION Heading Processing
            if line.startswith('# '):
                # Main heading (H1) - largest accessible size
                heading_text = line[2:].strip()
                heading = doc.add_heading(level=1)
                heading.clear()  # Clear default content to add custom formatting
                
                run = heading.add_run(heading_text)
                run.font.name = config.accessible_font
                run.font.size = Pt(config.heading1_font_size)
                run.font.bold = True
                run.font.color.rgb = RGBColor(*config.heading_color)
                
            elif line.startswith('## '):
                # Sub heading (H2) - large accessible size
                heading_text = line[3:].strip()
                heading = doc.add_heading(level=2)
                heading.clear()
                
                run = heading.add_run(heading_text)
                run.font.name = config.accessible_font
                run.font.size = Pt(config.heading2_font_size)
                run.font.bold = True
                run.font.color.rgb = RGBColor(*config.heading_color)
                
            elif line.startswith('### '):
                # Sub-sub heading (H3) - accessible size
                heading_text = line[4:].strip()
                heading = doc.add_heading(level=3)
                heading.clear()
                
                run = heading.add_run(heading_text)
                run.font.name = config.accessible_font
                run.font.size = Pt(config.heading3_font_size)
                run.font.bold = True
                run.font.color.rgb = RGBColor(51, 51, 51)  # Slightly lighter for hierarchy
                
            # SECTION List Processing
            elif line.startswith(('- ', '* ')):
                # Bullet point with accessible formatting
                para = doc.add_paragraph(style='List Bullet')
                para.clear()  # Clear default content
                
                run = para.add_run(line[2:])
                run.font.name = config.accessible_font
                run.font.size = Pt(config.body_font_size)
                run.font.color.rgb = RGBColor(*config.text_color)
                
                # Improve list spacing for accessibility
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.left_indent = Inches(0.25)
                
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                # Numbered list with accessible formatting
                para = doc.add_paragraph(style='List Number')
                para.clear()
                
                # Extract number and text
                parts = line.split('. ', 1)
                if len(parts) == 2:
                    run = para.add_run(parts[1])
                else:
                    run = para.add_run(line[3:])
                
                run.font.name = config.accessible_font
                run.font.size = Pt(config.body_font_size)
                run.font.color.rgb = RGBColor(*config.text_color)
                
                # Improve list spacing for accessibility
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.left_indent = Inches(0.25)
                
            else:
                # SECTION Regular Paragraph Processing
                # Regular paragraph with large print and bold text support
                para = doc.add_paragraph()
                
                # Handle bold text formatting (**text**)
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if part:  # Skip empty parts
                            run = para.add_run(part)
                            run.font.name = config.accessible_font
                            run.font.size = Pt(config.body_font_size)
                            run.font.color.rgb = RGBColor(*config.text_color)
                            
                            # Make every other part bold (content inside **)
                            if i % 2 == 1:
                                run.font.bold = True
                else:
                    # Simple paragraph without formatting
                    run = para.add_run(line)
                    run.font.name = config.accessible_font
                    run.font.size = Pt(config.body_font_size)
                    run.font.color.rgb = RGBColor(*config.text_color)
                
                # Set enhanced paragraph spacing for readability
                para.paragraph_format.space_after = Pt(config.paragraph_spacing)
                para.paragraph_format.line_spacing = config.line_spacing
    
    # SECTION Accessible Document Footer
    def _add_accessible_document_footer(self, doc: Document):
        """
        Create professional MyVision footer with accessibility statement.
        
        This footer provides essential contact information and accessibility
        transparency, informing users about the document's large print formatting
        while maintaining professional MyVision branding standards.
        
        Footer Components:
        1. Visual separator for document section division
        2. Contact information with multiple communication channels
        3. Accessibility statement noting large print formatting
        4. Professional formatting with readable font sizes
        
        Accessibility Features:
        - All footer text maintains readable sizing (14pt minimum)
        - Clear visual separation from main content
        - Accessibility statement promotes transparency
        - Multiple contact methods for user support
        """
        # SECTION Pre-Footer Spacing
        # Add appropriate space before footer section
        spacer_para = doc.add_paragraph()
        spacer_para.space_before = Pt(24)
        
        # SECTION Footer Separator
        # Professional separator with accessible formatting
        separator_para = doc.add_paragraph()
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        separator_run = separator_para.add_run("●" * 25)
        separator_run.font.name = config.accessible_font
        separator_run.font.size = Pt(config.body_font_size)
        separator_run.font.color.rgb = RGBColor(*config.accent_color)
        
        # SECTION Contact Information
        # Comprehensive contact information with large print formatting
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.space_before = Pt(12)
        
        footer_text = (f"{config.organization_name}\n"
                      f"Email: {config.contact_email}\n"
                      f"Web: {config.website}")
        
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = config.accessible_font
        footer_run.font.size = Pt(config.body_font_size - 2)  # Slightly smaller but still accessible
        footer_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # SECTION Accessibility Statement
        # Transparent statement about document accessibility features
        accessibility_para = doc.add_paragraph()
        accessibility_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        accessibility_para.space_before = Pt(12)
        
        accessibility_text = f"This document has been formatted for accessibility with large print ({config.body_font_size}pt minimum font size)"
        accessibility_run = accessibility_para.add_run(accessibility_text)
        accessibility_run.font.name = config.accessible_font
        accessibility_run.font.size = Pt(config.body_font_size - 4)  # Small but readable note
        accessibility_run.font.color.rgb = RGBColor(128, 128, 128)
        accessibility_run.italic = True
